

"""
SEO Optimizer Agent
- SEO polish + clean export to DOCX without markdown/HTML artifacts
- No diagrams, no code blocks, no schema or comments
- Validates references contain real URLs (no 'URL needed')
- Converts URLs to clickable hyperlinks in .docx
"""

import os
import re
import textwrap
from dotenv import load_dotenv
from langchain_openai import ChatOpenAI
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

load_dotenv()
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
llm = ChatOpenAI(model="gpt-4o-mini", temperature=0.3, openai_api_key=OPENAI_API_KEY)

INPUT_MD = "output/blog_edited.md"
OUT_MD   = "output/blog_seo.md"
OUT_DOCX = "ai-blog-automation/blog_final.docx"

# ---------- IO ----------

def read_markdown(path=INPUT_MD):
    if not os.path.exists(path):
        print("[!] blog_edited.md missing. Run editor_agent first.")
        return None
    with open(path, "r", encoding="utf-8") as f:
        return f.read()

# ---------- LLM SEO pass (no placeholders, no HTML/JS) ----------

def optimise_markdown(md: str) -> str:
    prompt = textwrap.dedent(f"""
<optimized_prompt>
<role>You are an expert SEO Optimizer and Content Formatter specializing in technical content. Your primary goal is to transform a given technical blog post into its final, SEO-optimized, and perfectly formatted version ready for publication. You must ensure maximum search engine visibility while strictly adhering to all specified formatting and content rules, preserving technical accuracy and readability.</role>
<instructions>
    <task>Take the provided technical blog post, apply comprehensive SEO optimization, and meticulously reformat it according to the strict output requirements. This includes optimizing all headings, ensuring clean Markdown, integrating FAQs, and managing references with real URLs.</task>
    <think_process>
        <step>1. Analyze the provided blog post and its core topic to identify primary, secondary, and long-tail keywords relevant to a technical audience.</step>
        <step>2. Craft an SEO-optimized title and meta description that are compelling and keyword-rich.</step>
        <step>3. **Rewrite all headings and subheadings (H1, H2, H3) to be highly SEO-optimized**, incorporating relevant keywords naturally while maintaining clarity and technical accuracy.</step>
        <step>4. Review the body content for natural keyword integration, ensuring density is appropriate and keywords enhance meaning. Suggest strategic placements for improved visibility.</step>
        <step>5. **Strictly ensure all content is presented in clean Markdown headings (#, ##, ###) and paragraphs. Remove any and all HTML tags, script, schema markup, comments, diagrams (if not text-based), code fences, or Mermaid syntax.**</step>
        <step>6. **Verify that any mentions of tools are integrated within paragraphs and are NOT presented as separate headings.**</step>
        <step>7. Identify opportunities for internal and external linking to relevant technical resources or other company blogs, ensuring all external links are to real, direct URLs.</step>
        <step>8. **Generate 2-3 relevant, concise, and SEO-friendly Frequently Asked Questions (FAQs) based on the blog post's content, to be placed at the end.**</step>
        <step>9. **For all references, ensure they include real, direct URLs. If a URL cannot be provided or is a placeholder (e.g., "URL needed", "TBD", empty links), remove that reference and its in-text marker entirely.**</step>
        <step>10. **Compile a 'References' section at the very end, listing every in-text citation [^n] with the exact format: Title (Year) -- URL.**</step>
        <step>11. Conduct a final review to ensure the entire output is polished, technically sound, and adheres to all formatting constraints.</step>
    </think_process>
    <output_requirements>
        <format>Clean Markdown only. No HTML, script, schema, comments, diagrams, code fences, or Mermaid syntax.</format>
        <headings_and_subheadings>All headings and subheadings must be SEO-optimized (H1, H2, H3).</headings_and_subheadings>
        <tool_mentions>Tools must be mentioned inside paragraphs, not as separate headings.</tool_mentions>
        <faqs>Include 2-3 relevant FAQs at the end of the blog post.</faqs>
        <references>
            <url_policy>All references must include real, direct URLs. No placeholders. Remove references without valid URLs.</url_policy>
            <format>A dedicated 'References' section at the very end. Each reference listed as: Title (Year) -- URL.</format>
        </references>
    </output_requirements>
</instructions>
<input_data>
    <blog_post>{md}</blog_post>
    <topic>{{BLOG_TOPIC}}</topic>
</input_data>
<output_format>
    <!-- The SEO Optimizer Agent will output the full, optimized blog post here, followed by FAQs and References. -->
    <optimized_blog_post>
        # [SEO Optimized Title]
        <!-- Meta Description (a short paragraph after title) -->
        [Optimized Meta Description Text just a short paragraph after title with no headings]

        ## [SEO Optimized H2 Heading]
        [Paragraph content with seo keyword integration. Tools mentioned within paragraphs.]

        ### [SEO Optimized H3 Heading]
        [Paragraph content.]

        ## [Another SEO Optimized H2 Heading]
        [Paragraph content with seo keyword integration. Internal/external links integrated.]

        ### [Another SEO Optimized H3 Heading]
        [Paragraph content.]

        <!-- ... Continue with the rest of the blog post content ... -->

        ## FAQ's

        ### [SEO Optimized FAQ Question 1]?
        [Answer to FAQ 1.]

        ### [SEO Optimized FAQ Question 2]?
        [Answer to FAQ 2.]

        ### [SEO Optimized FAQ Question 3]?
        [Answer to FAQ 3.]

        ## References

        1: Title of Reference 1 (Year) -- https://www.example.com/reference1
        2: Title of Reference 2 (Year) -- https://www.example.com/reference2
        <!-- ... List all in-text citations here ... -->
    </optimized_blog_post>
</output_format>
</optimized_prompt>
""")
    return llm.invoke(prompt).content

# ---------- Validation: ensure references have real URLs ----------

def extract_references(md_text: str):
    """
    Grab the References section (simple heuristic): lines after a heading that starts with '## References' or '### References'
    Returns a list of reference lines.
    """
    refs = []
    lines = md_text.splitlines()
    in_refs = False
    for line in lines:
        if re.match(r"^\s*##+\s+References\s*$", line, flags=re.IGNORECASE):
            in_refs = True
            continue
        if in_refs:
            # Stop when hitting a new H2/H3 or end
            if re.match(r"^\s*##+\s+", line):
                break
            if line.strip():
                refs.append(line.strip())
    return refs

def references_have_valid_urls(ref_lines):
    """
    Checks that each reference line includes at least one http(s) URL
    and no placeholders like 'URL needed'.
    """
    bad = []
    for idx, line in enumerate(ref_lines, start=1):
        has_placeholder = re.search(r"URL\s*needed|TBD|PLACEHOLDER", line, flags=re.IGNORECASE)
        has_url = re.search(r"https?://[^\s)]+", line)
        if has_placeholder or not has_url:
            bad.append((idx, line))
    return bad  # list of (index, line) with issues

def remove_bad_references(md_text: str, bad_refs):
    """
    Remove lines from the References section that are in bad_refs.
    """
    if not bad_refs:
        return md_text
    lines = md_text.splitlines()
    refs_start = None
    refs_end = None
    # Find start and end of References section
    for i, line in enumerate(lines):
        if re.match(r"^\s*##+\s+References\s*$", line, flags=re.IGNORECASE):
            refs_start = i
            continue
        if refs_start is not None and re.match(r"^\s*##+\s+", line):
            refs_end = i
            break
    if refs_start is None:
        return md_text
    refs_end = refs_end or len(lines)
    # Remove bad reference lines
    ref_lines = lines[refs_start+1:refs_end]
    bad_lines = set(line for _, line in bad_refs)
    new_ref_lines = [line for line in ref_lines if line.strip() and line.strip() not in bad_lines]
    # Rebuild lines
    new_lines = lines[:refs_start+1] + new_ref_lines + lines[refs_end:]
    return "\n".join(new_lines)

# ---------- DOCX: strip markdown/HTML artifacts & add hyperlinks ----------

def add_hyperlink(paragraph, url_text):
    """
    Insert a clickable hyperlink run into a paragraph (python-docx low-level).
    url_text is both display text and the target.
    """
    part = paragraph.part
    r_id = part.relate_to(url_text, reltype="http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')
    rPr.append(color)

    new_run.append(rPr)
    t = OxmlElement('w:t')
    t.text = url_text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

def strip_markdown_and_export_docx(md_text: str, docx_path: str):
    """
    Convert Markdown to Word:
    - Remove comments, scripts, code fences
    - Convert headings to bold paragraphs (no # or **)
    - Strip inline markdown emphasis markers ** * _
    - Convert plain URLs in text to clickable hyperlinks
    - Keep bullet/numbered lists
    - Bold text surrounded by double asterisks (**like this**)
    """
    cleaned = re.sub(r"<!--.*?-->", "", md_text, flags=re.DOTALL)
    cleaned = re.sub(r"<script.*?>.*?</script>", "", cleaned, flags=re.DOTALL)
    cleaned = re.sub(r"```.*?```", "", cleaned, flags=re.DOTALL)
    cleaned = cleaned.replace("\r\n", "\n")

    doc = Document()

    def add_heading(text, level):
        p = doc.add_paragraph()
        last = 0
        for m in re.finditer(r"\*\*(.+?)\*\*", text):
            if m.start() > last:
                p.add_run(text[last:m.start()])
            p.add_run(m.group(1)).bold = True
            last = m.end()
        if last < len(text):
            p.add_run(text[last:])
        for run in p.runs:
            run.bold = True
            if level == 1:
                run.font.size = Pt(16)
            elif level == 2:
                run.font.size = Pt(14)
            else:
                run.font.size = Pt(12)

    url_pattern = re.compile(r"(https?://[^\s)]+)")

    def add_bolded_paragraph(text, style=None):
        p = doc.add_paragraph("" if style else None, style=style)
        last = 0
        for m in re.finditer(r"\*\*(.+?)\*\*", text):
            if m.start() > last:
                p.add_run(text[last:m.start()])
            p.add_run(m.group(1)).bold = True
            last = m.end()
        if last < len(text):
            p.add_run(text[last:])
        for m in url_pattern.findall(text):
            add_hyperlink(p, m)
        return p

    for raw_line in cleaned.split("\n"):
        line = raw_line.rstrip()
        if not line.strip():
            doc.add_paragraph("")
            continue

        if re.match(r"^\s*#\s+", line):
            add_heading(re.sub(r"^\s*#\s+", "", line), 1)
            continue
        if re.match(r"^\s*##\s+", line):
            add_heading(re.sub(r"^\s*##\s+", "", line), 2)
            continue
        if re.match(r"^\s*###\s+", line):
            add_heading(re.sub(r"^\s*###\s+", "", line), 3)
            continue

        if re.match(r"^\s*-\s+", line):
            text = re.sub(r"^\s*-\s+", "", line)
            add_bolded_paragraph(text, style="List Bullet")
            continue

        if re.match(r"^\s*\d+\.\s+", line):
            text = re.sub(r"^\s*\d+\.\s+", "", line)
            add_bolded_paragraph(text, style="List Number")
            continue

        add_bolded_paragraph(line)

    doc.save(docx_path)

# ---------- Orchestration ----------

def run():
    original_md = read_markdown()
    if not original_md:
        return

    print("[~] Optimizing markdown (no placeholders, no HTML/JS)…")
    optimised_md = optimise_markdown(original_md)

    # Validate references
    refs = extract_references(optimised_md)
    bad_refs = references_have_valid_urls(refs)
    if bad_refs:
        print("Found references without real URLs or with placeholders:")
        for idx, line in bad_refs:
            print(f"   ({idx}) {line}")
        print("\nRemoving bad references and continuing export…")
        # Remove bad references and continue
        optimised_md = remove_bad_references(optimised_md, bad_refs)

    # Save cleaned markdown
    os.makedirs("output", exist_ok=True)
    with open(OUT_MD, "w", encoding="utf-8") as f:
        f.write(optimised_md)
    print(f"[✓] SEO-optimized markdown saved to {OUT_MD}")

    # Export to DOCX with clickable links
    strip_markdown_and_export_docx(optimised_md, OUT_DOCX)
    print(f"Word doc saved to {OUT_DOCX}")

if __name__ == "__main__":
    run()